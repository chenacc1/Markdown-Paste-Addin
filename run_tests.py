#!/usr/bin/env python3
"""Comprehensive test suite for MarkdownPasteAddin v3.0"""
import sys, os, json, time, tempfile, subprocess

os.chdir(os.path.dirname(os.path.abspath(__file__)))

print('=' * 60)
print('  MarkdownPasteAddin v3.0 -- Comprehensive Test Suite')
print('=' * 60)
print(f'  Time: {time.strftime("%Y-%m-%d %H:%M:%S")}')
print()

results = {}
passed = 0
failed = 0

def test(name, fn):
    global passed, failed
    try:
        fn()
        results[name] = 'PASS'
        passed += 1
        print(f'  [PASS] {name}')
    except Exception as e:
        results[name] = f'FAIL: {e}'
        failed += 1
        print(f'  [FAIL] {name}: {e}')

# 1. Module imports
def t_imports():
    from md2docx_lib import (parse_markdown, parse_html, build_docx, DocumentBuilder)
    from md2docx_lib.parser_math import extract_math, latex_to_omml, MathBlock
    from md2docx_lib.inline_processor import process_inline_paragraph, clean_text, _tokenize
    from md2docx_lib.report_standards import PRESETS, apply_report_standard
    from md2docx_lib.builder_numbering import NumberingTracker
    from md2docx_lib.builder_toc import add_toc
    from md2docx_lib.formatter import FORMAT, format_document, detect_paragraph_type
    from md2docx_lib.template import load_template
    assert len(PRESETS) == 3
test('Module imports (15 modules)', t_imports)

# 2. Markdown parsing
def t_markdown_parsing():
    from md2docx_lib import parse_markdown
    md = open('tests/MarkdownPasteAddin.Tests/TestData/sample-mixed.md', encoding='utf-8').read()
    chunks = parse_markdown(md)
    types = [c['type'] for c in chunks]
    assert 'heading' in types
    assert 'table' in types
    assert 'mermaid' in types
    assert len(chunks) >= 5
test('Markdown parsing (sample-mixed.md)', t_markdown_parsing)

# 3. All chunk types
def t_all_chunk_types():
    from md2docx_lib import parse_markdown
    md = "# H1\n## H2\ntext\n|A|B|\n|---|---|\n|1|2|\n```python\ncode\n```\n```mermaid\nA-->B\n```\n- [ ] task\n- [x] done\n> quote\n---\n![img](http://x.com/p.png)\n"
    chunks = parse_markdown(md)
    types = set(c['type'] for c in chunks)
    expected = {'heading','text','table','code','mermaid','task_list','blockquote','hr','image'}
    missing = expected - types
    assert not missing, f'Missing: {missing}'
test('All 9 chunk types', t_all_chunk_types)

# 4. Table alignments
def t_table_parsing():
    from md2docx_lib import parse_markdown
    md = '|L|C|R|\n|:--|:--:|--:|\n|a|b|c|'
    chunks = parse_markdown(md)
    t = chunks[0]
    assert t['alignments'] == ['left','center','right']
    assert t['headers'] == ['L','C','R']
test('Table alignment parsing', t_table_parsing)

# 5. Math extraction
def t_math_extraction():
    from md2docx_lib.parser_math import extract_math
    blocks = extract_math('Hello $x^2$ world $$y^3$$ end')
    assert len(blocks) == 2
    assert blocks[0].display == False and blocks[0].text == 'x^2'
    assert blocks[1].display == True and blocks[1].text == 'y^3'
test('Math formula extraction', t_math_extraction)

# 6. LaTeX to OMML
def t_latex_to_omml():
    from md2docx_lib.parser_math import latex_to_omml
    xml = latex_to_omml('x^2 + y^2')
    assert '<m:' in xml
    xml2 = latex_to_omml(r'\frac{a}{b}')
    assert '<m:f>' in xml2
    xml3 = latex_to_omml(r'\sqrt{x}')
    assert '<m:rad>' in xml3
test('LaTeX to OMML conversion', t_latex_to_omml)

# 7. Inline markdown
def t_inline_processing():
    from md2docx_lib.inline_processor import _tokenize, clean_text
    tokens = _tokenize('Hello **bold** and *italic* with `code` end.')
    tok_types = [t for t,_ in tokens]
    assert 'bold' in tok_types
    assert 'italic' in tok_types
    assert 'code' in tok_types
    cleaned = clean_text('**bold** *italic* `code` ~~strike~~')
    assert '**' not in cleaned
    assert '`' not in cleaned
    assert '~~' not in cleaned
test('Inline markdown processing', t_inline_processing)

# 8. Report standards
def t_report_standards():
    from md2docx_lib.report_standards import PRESETS
    for k in ['govt', 'academic', 'business']:
        assert k in PRESETS
        p = PRESETS[k]
        assert len(p.styles) >= 8, f'{k}: {len(p.styles)} styles'
        assert p.cover.enabled
        assert p.page.margin_top is not None
test('Report standards (3 presets with 8 styles each)', t_report_standards)

# 9. Numbering tracker
def t_numbering():
    from md2docx_lib.builder_numbering import NumberingTracker
    nt = NumberingTracker()
    assert nt.next_figure() == '图 1'
    assert nt.next_figure('desc') == '图 2：desc'
    assert nt.next_table() == '表 1'
    assert nt.next_table('data') == '表 2：data'
test('Figure/table auto-numbering', t_numbering)

# 10. Full pipeline
def t_full_pipeline():
    from md2docx_lib import parse_markdown
    from md2docx_lib.builder_docx import DocumentBuilder
    from md2docx_lib.report_standards import apply_report_standard
    from docx import Document
    md = "# Report\n## Section\n**Bold** and *italic* text.\n| A | B |\n|---|---|\n| 1 | 2 |\n```python\ndef hello(): pass\n```"
    path = os.path.join(tempfile.gettempdir(), 'test_report_pipeline.docx')
    chunks = parse_markdown(md)
    builder = DocumentBuilder(auto_captions=True, add_toc=True)
    builder.build(chunks, path)
    assert os.path.exists(path)
    size1 = os.path.getsize(path)
    assert size1 > 5000
    doc = Document(path)
    apply_report_standard(doc, preset='business', title='Test Report')
    doc.save(path)
    size2 = os.path.getsize(path)
    assert size2 > size1
    os.unlink(path)
test('Full pipeline (parse -> build -> format)', t_full_pipeline)

# 11. HTML parsing
def t_html_parsing():
    from md2docx_lib.parser_html import parse_html
    html = '<h1>Title</h1><p>Text</p><table><tr><th>A</th></tr><tr><td>1</td></tr></table>'
    chunks = parse_html(html)
    types = [c['type'] for c in chunks]
    assert 'table' in types
test('HTML parsing', t_html_parsing)

# 12. Bridge server
def t_bridge_health():
    import urllib.request
    proc = subprocess.Popen([sys.executable, 'bridge_server.py', '--port', '19878'],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    time.sleep(2)
    try:
        resp = urllib.request.urlopen('http://localhost:19878/api/health', timeout=3)
        data = json.loads(resp.read())
        assert data['status'] == 'ok'
        assert 'version' in data
    finally:
        proc.terminate()
        proc.wait()
test('Bridge server health check', t_bridge_health)

# 13. Python syntax check
def t_syntax():
    import py_compile
    lib_dir = 'md2docx_lib'
    for f in os.listdir(lib_dir):
        if f.endswith('.py'):
            py_compile.compile(os.path.join(lib_dir, f), doraise=True)
    for f in ['md2docx.py', 'format_docx.py', 'gui_app.py', 'bridge_server.py',
              'batch_convert.py', 'watch_convert.py', 'deepseek_api.py']:
        py_compile.compile(f, doraise=True)
test('Python syntax (all 22 .py files)', t_syntax)

# 14. docx content verification
def t_docx_content():
    from md2docx_lib import parse_markdown
    from md2docx_lib.builder_docx import DocumentBuilder
    from docx import Document
    md = "# Test\n\nParagraph with **bold** text.\n\n|X|Y|\n|---|---|\n|a|b|"
    path = os.path.join(tempfile.gettempdir(), 'test_content.docx')
    chunks = parse_markdown(md)
    DocumentBuilder(show_progress=False).build(chunks, path)
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = doc.tables
    assert any('Test' in t for t in paragraphs), f'Title not found in {paragraphs}'
    assert any('bold' in t.lower() for t in paragraphs)
    assert len(tables) == 1
    assert tables[0].cell(0,0).text == 'X'
    os.unlink(path)
test('Generated docx content verification', t_docx_content)

# Summary
print()
print('=' * 60)
print(f'  Results: {passed} PASS, {failed} FAIL, {passed+failed} total')
if failed > 0:
    print(f'  Failed: {[k for k,v in results.items() if "FAIL" in str(v)]}')
print('=' * 60)
sys.exit(0 if failed == 0 else 1)
