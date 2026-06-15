"""
Professional report format standards for Word documents.

Supports three presets:
  - 'govt'    : GB/T 9704-2012 党政机关公文格式 (中国政府公文标准)
  - 'academic': 学术论文格式 (高校学位论文通用格式)
  - 'business': 通用商务报告格式 (企业报告/技术文档推荐格式)

Usage:
  from md2docx_lib.report_standards import apply_report_standard
  apply_report_standard(doc, preset='business')

  # Get a preset's config for custom use:
  from md2docx_lib.report_standards import PRESETS
  preset = PRESETS['business']
"""

from dataclasses import dataclass, field
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Inches, RGBColor, Emu


# ═══════════════════════════════════════════════════════════════════════════════
# Format preset data structures
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class FontSpec:
    """Font specification with CJK support."""
    name: str               # Latin font
    east_asian: str = ""    # CJK font (falls back to name if empty)
    size: Pt = Pt(12)
    bold: bool = False
    italic: bool = False
    color: RGBColor = RGBColor(0, 0, 0)

@dataclass
class ParaSpec:
    """Paragraph formatting specification."""
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT
    space_before: Pt = Pt(0)
    space_after: Pt = Pt(0)
    line_spacing: float = 1.5
    line_spacing_rule: int = WD_LINE_SPACING.MULTIPLE
    first_line_indent: Cm | None = None   # None = no indent
    outline_level: int = -1               # -1 = not set, 0-8 = heading level
    keep_with_next: bool = False

@dataclass
class StyleSpec:
    """Combined font + paragraph style for a semantic element."""
    name: str
    font: FontSpec
    paragraph: ParaSpec

@dataclass
class PageSpec:
    """Page layout specification."""
    paper_width: Cm = Cm(21.0)      # A4 width
    paper_height: Cm = Cm(29.7)     # A4 height
    margin_top: Cm = Cm(2.54)
    margin_bottom: Cm = Cm(2.54)
    margin_left: Cm = Cm(3.17)
    margin_right: Cm = Cm(3.17)
    header_distance: Cm = Cm(1.5)
    footer_distance: Cm = Cm(1.75)

@dataclass
class CoverSpec:
    """Cover page specification."""
    enabled: bool = True
    title_font: FontSpec = field(default_factory=lambda: FontSpec("黑体", "黑体", Pt(22), True))
    subtitle_font: FontSpec = field(default_factory=lambda: FontSpec("黑体", "黑体", Pt(16), False))
    info_font: FontSpec = field(default_factory=lambda: FontSpec("宋体", "宋体", Pt(14), False))
    show_date: bool = True
    show_author: bool = True
    show_org: bool = True

@dataclass
class HeaderFooterSpec:
    """Header/footer specification."""
    show_page_number: bool = True
    page_number_align: int = WD_ALIGN_PARAGRAPH.CENTER
    page_number_format: str = "第{page}页 / 共{total}页"
    show_header_line: bool = True    # thin line below header
    header_text: str = ""            # empty = use document title
    footer_text: str = ""


@dataclass
class ReportPreset:
    """Complete report format preset."""
    name: str
    description: str
    page: PageSpec
    cover: CoverSpec
    header_footer: HeaderFooterSpec
    styles: dict[str, StyleSpec]     # key: 'h1','h2','h3','body','fig_caption','tbl_caption','code','quote'
    title_page_style: StyleSpec | None = None  # special style for document title on first page


# ═══════════════════════════════════════════════════════════════════════════════
# Presets
# ═══════════════════════════════════════════════════════════════════════════════

PRESETS: dict[str, ReportPreset] = {}

# ── 中国政府公文标准 (GB/T 9704-2012) ──────────────────────────────────────────

PRESETS["govt"] = ReportPreset(
    name="党政机关公文格式 (GB/T 9704-2012)",
    description="中国政府机关正式公文标准格式。适用于红头文件、通知、报告等正式公文。",
    page=PageSpec(
        paper_width=Cm(21.0), paper_height=Cm(29.7),
        margin_top=Cm(3.7), margin_bottom=Cm(3.5),
        margin_left=Cm(2.8), margin_right=Cm(2.6),
    ),
    cover=CoverSpec(
        enabled=True,
        title_font=FontSpec("方正小标宋简体", "方正小标宋简体", Pt(22), True),
        subtitle_font=FontSpec("楷体", "楷体", Pt(16), False),
        info_font=FontSpec("仿宋", "仿宋", Pt(16), False),
    ),
    header_footer=HeaderFooterSpec(
        show_page_number=True,
        page_number_align=WD_ALIGN_PARAGRAPH.CENTER,
        page_number_format="— {page} —",
        show_header_line=False,
    ),
    styles={
        "h1": StyleSpec("一级标题",
            FontSpec("黑体", "黑体", Pt(16), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(20), Pt(10), 1.5, outline_level=1)),
        "h2": StyleSpec("二级标题",
            FontSpec("楷体", "楷体", Pt(16), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(16), Pt(8), 1.5, outline_level=2)),
        "h3": StyleSpec("三级标题",
            FontSpec("仿宋", "仿宋", Pt(16), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(12), Pt(6), 1.5, outline_level=3)),
        "body": StyleSpec("正文",
            FontSpec("仿宋", "仿宋", Pt(16), False),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0), 1.5,
                     first_line_indent=Cm(0.85))),
        "fig_caption": StyleSpec("图题",
            FontSpec("楷体", "楷体", Pt(14), False),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(12), 1.5)),
        "tbl_caption": StyleSpec("表题",
            FontSpec("楷体", "楷体", Pt(14), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(6), 1.5)),
        "code": StyleSpec("代码",
            FontSpec("Courier New", "宋体", Pt(11), False, color=RGBColor(50,50,50)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0), 1.0)),
        "quote": StyleSpec("引用",
            FontSpec("楷体", "楷体", Pt(14), False, color=RGBColor(80,80,80)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(6), 1.5,
                     first_line_indent=Cm(0.85))),
    },
)

# ── 学术论文格式 ────────────────────────────────────────────────────────────────

PRESETS["academic"] = ReportPreset(
    name="学术论文格式",
    description="高校学位论文/学术期刊通用格式。适用于毕业论文、学术文章、研究报告。",
    page=PageSpec(
        margin_top=Cm(2.54), margin_bottom=Cm(2.54),
        margin_left=Cm(3.17), margin_right=Cm(3.17),
    ),
    cover=CoverSpec(
        enabled=True,
        title_font=FontSpec("黑体", "黑体", Pt(22), True),
        subtitle_font=FontSpec("黑体", "黑体", Pt(16), False),
        info_font=FontSpec("宋体", "宋体", Pt(14), False),
    ),
    header_footer=HeaderFooterSpec(
        show_page_number=True,
        page_number_align=WD_ALIGN_PARAGRAPH.CENTER,
        page_number_format="{page}",
        show_header_line=True,
        header_text="",
    ),
    styles={
        "h1": StyleSpec("一级标题",
            FontSpec("黑体", "黑体", Pt(16), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(24), Pt(12), 1.5, outline_level=1)),
        "h2": StyleSpec("二级标题",
            FontSpec("黑体", "黑体", Pt(14), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(18), Pt(8), 1.5, outline_level=2)),
        "h3": StyleSpec("三级标题",
            FontSpec("黑体", "黑体", Pt(12), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(12), Pt(6), 1.5, outline_level=3)),
        "body": StyleSpec("正文",
            FontSpec("宋体", "宋体", Pt(12), False),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0), 1.5,
                     first_line_indent=Cm(0.74))),
        "fig_caption": StyleSpec("图题",
            FontSpec("宋体", "宋体", Pt(10), False),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(12), 1.5)),
        "tbl_caption": StyleSpec("表题",
            FontSpec("宋体", "宋体", Pt(10), True),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(6), 1.5)),
        "code": StyleSpec("代码",
            FontSpec("Consolas", "宋体", Pt(9), False, color=RGBColor(50,50,50)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(4), Pt(4), 1.0)),
        "quote": StyleSpec("引用",
            FontSpec("楷体", "楷体", Pt(12), False, color=RGBColor(80,80,80)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(6), 1.5,
                     first_line_indent=Cm(0.74))),
    },
)

# ── 通用商务报告格式 ────────────────────────────────────────────────────────────

PRESETS["business"] = ReportPreset(
    name="通用商务报告格式",
    description="企业报告、技术文档、产品手册推荐格式。兼容性最好，适用于大多数场景。",
    page=PageSpec(
        margin_top=Cm(2.54), margin_bottom=Cm(2.54),
        margin_left=Cm(2.54), margin_right=Cm(2.54),
    ),
    cover=CoverSpec(
        enabled=True,
        title_font=FontSpec("Arial", "黑体", Pt(26), True, color=RGBColor(30,60,120)),
        subtitle_font=FontSpec("Arial", "黑体", Pt(16), False, color=RGBColor(60,90,150)),
        info_font=FontSpec("Calibri", "宋体", Pt(12), False, color=RGBColor(100,100,100)),
        show_date=True,
        show_author=True,
        show_org=True,
    ),
    header_footer=HeaderFooterSpec(
        show_page_number=True,
        page_number_align=WD_ALIGN_PARAGRAPH.RIGHT,
        page_number_format="{page} / {total}",
        show_header_line=True,
        header_text="",
    ),
    styles={
        "h1": StyleSpec("一级标题",
            FontSpec("Arial", "黑体", Pt(18), True, color=RGBColor(30,60,120)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(24), Pt(12), 1.15, outline_level=1)),
        "h2": StyleSpec("二级标题",
            FontSpec("Arial", "黑体", Pt(15), True, color=RGBColor(45,75,135)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(18), Pt(8), 1.15, outline_level=2)),
        "h3": StyleSpec("三级标题",
            FontSpec("Arial", "黑体", Pt(13), True, color=RGBColor(60,90,150)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(12), Pt(6), 1.15, outline_level=3)),
        "body": StyleSpec("正文",
            FontSpec("Calibri", "宋体", Pt(11), False),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(6), 1.15,
                     first_line_indent=Cm(0.0))),
        "fig_caption": StyleSpec("图题",
            FontSpec("Calibri", "宋体", Pt(9), False, color=RGBColor(80,80,80)),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(4), Pt(10), 1.15)),
        "tbl_caption": StyleSpec("表题",
            FontSpec("Calibri", "宋体", Pt(9), True, color=RGBColor(60,60,60)),
            ParaSpec(WD_ALIGN_PARAGRAPH.CENTER, Pt(10), Pt(4), 1.15)),
        "code": StyleSpec("代码",
            FontSpec("Consolas", "Consolas", Pt(9), False, color=RGBColor(50,55,60)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(0), Pt(0), 1.0)),
        "quote": StyleSpec("引用",
            FontSpec("Calibri", "宋体", Pt(10.5), True, color=RGBColor(100,100,110)),
            ParaSpec(WD_ALIGN_PARAGRAPH.LEFT, Pt(8), Pt(8), 1.15,
                     first_line_indent=Cm(0.0))),
    },
)


# ═══════════════════════════════════════════════════════════════════════════════
# Page setup
# ═══════════════════════════════════════════════════════════════════════════════

def apply_page_setup(doc: Document, preset: ReportPreset):
    """Apply page dimensions and margins from preset to all sections."""
    page = preset.page
    for section in doc.sections:
        section.page_width = page.paper_width
        section.page_height = page.paper_height
        section.top_margin = page.margin_top
        section.bottom_margin = page.margin_bottom
        section.left_margin = page.margin_left
        section.right_margin = page.margin_right
        section.header_distance = page.header_distance
        section.footer_distance = page.footer_distance


# ═══════════════════════════════════════════════════════════════════════════════
# Style application
# ═══════════════════════════════════════════════════════════════════════════════

_WD_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: 0,
    WD_ALIGN_PARAGRAPH.CENTER: 1,
    WD_ALIGN_PARAGRAPH.RIGHT: 2,
    WD_ALIGN_PARAGRAPH.JUSTIFY: 3,
}


def apply_style_to_paragraph(para, style: StyleSpec):
    """Apply a StyleSpec (font + paragraph formatting) to a paragraph."""
    _apply_para_format(para, style.paragraph)
    for run in para.runs:
        _apply_font(run, style.font)
    if not para.runs and para.text.strip():
        run = para.add_run(para.text)
        _apply_font(run, style.font)


def _apply_para_format(para, ps: ParaSpec):
    pf = para.paragraph_format
    pf.alignment = ps.alignment
    if ps.space_before is not None:
        pf.space_before = ps.space_before
    if ps.space_after is not None:
        pf.space_after = ps.space_after
    if ps.line_spacing is not None:
        pf.line_spacing = ps.line_spacing
    pf.line_spacing_rule = ps.line_spacing_rule
    if ps.first_line_indent is not None:
        pf.first_line_indent = ps.first_line_indent
    if ps.outline_level >= 0:
        _set_outline_level(para, ps.outline_level)
    if ps.keep_with_next:
        pf.keep_with_next = True


def _apply_font(run, fs: FontSpec):
    run.font.name = fs.name
    run.font.size = fs.size
    run.font.bold = fs.bold
    run.font.italic = fs.italic
    if fs.color:
        run.font.color.rgb = fs.color

    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    ea = fs.east_asian or fs.name
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:ascii"), fs.name)
    rFonts.set(qn("w:hAnsi"), fs.name)


def _set_outline_level(para, level: int):
    pPr = para._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(existing)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(level))
    pPr.append(el)


# ═══════════════════════════════════════════════════════════════════════════════
# Cover page
# ═══════════════════════════════════════════════════════════════════════════════

def add_cover_page(doc: Document, preset: ReportPreset, title: str = "",
                   subtitle: str = "", author: str = "", org: str = "",
                   date_str: str = ""):
    """Insert a cover page at the beginning of the document."""
    cover = preset.cover
    if not cover.enabled:
        return

    # Only add cover if we have at least a title
    if not title:
        return

    # Top spacers
    for _ in range(4):
        para = doc.paragraphs[0].insert_paragraph_before("")
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(60)

    # Title
    title_para = doc.paragraphs[0].insert_paragraph_before(title)
    _apply_para_format(title_para, ParaSpec(
        WD_ALIGN_PARAGRAPH.CENTER, Pt(60), Pt(18), 1.5, outline_level=-1))
    run = title_para.runs[0] if title_para.runs else title_para.add_run(title)
    _apply_font(run, cover.title_font)

    # Subtitle (only if present)
    if subtitle:
        sub_para = doc.paragraphs[0].insert_paragraph_before(subtitle)
        _apply_para_format(sub_para, ParaSpec(
            WD_ALIGN_PARAGRAPH.CENTER, Pt(0), Pt(12), 1.5))
        run = sub_para.runs[0] if sub_para.runs else sub_para.add_run(subtitle)
        _apply_font(run, cover.subtitle_font)

    # Decorative line (only if we have author/org/date)
    has_info = (author and cover.show_author) or (org and cover.show_org) or cover.show_date
    if has_info:
        line_para = doc.paragraphs[0].insert_paragraph_before("─" * 36)
        _apply_para_format(line_para, ParaSpec(
            WD_ALIGN_PARAGRAPH.CENTER, Pt(18), Pt(12), 1.0))
        if line_para.runs:
            line_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    # Info lines
    if author and cover.show_author:
        _add_cover_info(doc, f"作者：{author}", cover)
    if org and cover.show_org:
        _add_cover_info(doc, f"单位：{org}", cover)
    if date_str and cover.show_date:
        _add_cover_info(doc, f"日期：{date_str}", cover)
    if not date_str and cover.show_date:
        from datetime import date
        _add_cover_info(doc, f"日期：{date.today().strftime('%Y年%m月%d日')}", cover)

    # Page break after cover
    doc.paragraphs[0].insert_paragraph_before("")


def _add_cover_info(doc: Document, text: str, cover: CoverSpec):
    para = doc.paragraphs[0].insert_paragraph_before(text)
    _apply_para_format(para, ParaSpec(
        WD_ALIGN_PARAGRAPH.CENTER, Pt(0), Pt(6), 1.5))
    run = para.runs[0] if para.runs else para.add_run(text)
    _apply_font(run, cover.info_font)


# ═══════════════════════════════════════════════════════════════════════════════
# Header / Footer
# ═══════════════════════════════════════════════════════════════════════════════

def add_header_footer(doc: Document, preset: ReportPreset, title: str = ""):
    """Add header and footer to all sections."""
    hf = preset.header_footer
    for section in doc.sections:
        # Footer with page number
        if hf.show_page_number:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = hf.page_number_align
            footer_para.paragraph_format.space_before = Pt(0)
            footer_para.paragraph_format.space_after = Pt(0)

            # Add page number field
            run = footer_para.add_run()
            _add_page_number_field(run, hf.page_number_format)

            run.font.size = Pt(9)
            run.font.name = "Calibri"

        # Header
        if hf.show_header_line or hf.header_text:
            header = section.header
            header.is_linked_to_previous = False
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if hf.header_text or title:
                header_para.text = hf.header_text or title
                header_para.runs[0].font.size = Pt(9) if header_para.runs else None

            # Add bottom border to header
            if hf.show_header_line:
                _add_paragraph_border(header_para, "bottom", "999999", 4)


def _add_page_number_field(run, fmt: str = "{page}"):
    """Insert Word PAGE / NUMPAGES fields."""
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_sep)

    run2_elem = OxmlElement("w:r")
    run2_text = OxmlElement("w:t")
    run2_text.text = "1"
    run2_elem.append(run2_text)
    run._r.addnext(run2_elem)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.addnext(fldChar_end)


def _add_paragraph_border(para, side: str, color: str, sz: int):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    pBdr.append(border)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# Main API
# ═══════════════════════════════════════════════════════════════════════════════

def apply_report_standard(doc: Document, preset: str = "business",
                          title: str = "", subtitle: str = "",
                          author: str = "", org: str = ""):
    """Apply a complete report format standard to an existing document.

    Args:
        doc: python-docx Document object
        preset: 'govt' | 'academic' | 'business'
        title: Document title (for cover and header)
        subtitle: Document subtitle
        author: Author name
        org: Organization name
    """
    if preset not in PRESETS:
        raise ValueError(f"Unknown preset '{preset}'. Available: {list(PRESETS.keys())}")

    p = PRESETS[preset]
    date_str = ""  # use today's date by default

    # 1. Page setup
    apply_page_setup(doc, p)

    # 2. Cover page
    add_cover_page(doc, p, title=title, subtitle=subtitle,
                   author=author, org=org, date_str=date_str)

    # 3. Header/footer
    add_header_footer(doc, p, title=title)

    # 4. Style all content paragraphs
    format_all_paragraphs(doc, p)


def format_all_paragraphs(doc: Document, preset: ReportPreset):
    """Apply style specifications to all paragraphs in a document."""
    styles = preset.styles
    import re

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Skip paragraphs that contain OMML math elements
        math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        if para._element.find(f".//{{{math_ns}}}oMath") is not None:
            continue
        if para._element.find(f".//{{{math_ns}}}oMathPara") is not None:
            continue

        # Detect paragraph type
        style_name = _detect_semantic_type(para, text)

        if style_name and style_name in styles:
            apply_style_to_paragraph(para, styles[style_name])
        else:
            apply_style_to_paragraph(para, styles.get("body", styles["body"]))


def _detect_semantic_type(para, text: str) -> str:
    """Detect semantic type from paragraph content."""
    import re

    # Already styled as heading
    sn = (para.style.name if para.style else "").lower()
    if "heading 1" in sn: return "h1"
    if "heading 2" in sn: return "h2"
    if "heading 3" in sn: return "h3"

    # Content-based detection
    if re.match(r"^(图|Fig\.?|Figure)\s*\d+", text, re.IGNORECASE):
        return "fig_caption"
    if re.match(r"^(表|Table)\s*\d+", text, re.IGNORECASE):
        return "tbl_caption"

    return "body"


def list_presets():
    """List available presets with descriptions."""
    for key, p in PRESETS.items():
        print(f"  [{key}] {p.name}")
        print(f"       {p.description}")
        print()
