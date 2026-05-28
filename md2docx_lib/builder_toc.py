"""Word TOC (Table of Contents) field code insertion."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def add_toc(doc: Document, title: str = "目录"):
    """Insert a Table of Contents at the current document position.

    Uses Word's TOC field code which gets rendered when the document
    is opened in Word and the user right-clicks → Update Field.
    """
    # Add title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    para.paragraph_format.space_after = Pt(12)

    # Insert TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # TOC field: fldChar begin
    fldChar_begin = _make_element("w:fldChar", {"w:fldCharType": "begin"})
    run._r.append(fldChar_begin)

    # instrText: TOC instructions
    instrText = _make_element("w:instrText")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    # fldChar separate
    fldChar_separate = _make_element("w:fldChar", {"w:fldCharType": "separate"})
    run._r.append(fldChar_separate)

    # Placeholder text
    run2 = paragraph.add_run('[ 请在 Word 中右键此处 → 更新域 以生成目录 ]')
    run2.font.color.rgb = None  # keep default
    run2.font.size = Pt(10)
    run2.font.italic = True

    # fldChar end
    run3 = paragraph.add_run()
    fldChar_end = _make_element("w:fldChar", {"w:fldCharType": "end"})
    run3._r.append(fldChar_end)

    # Page break after TOC
    doc.add_page_break()


def _make_element(tag: str, attrs: dict = None):
    """Create an lxml element with w: namespace."""
    from lxml import etree
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    el = etree.SubElement(etree.Element("{" + ns + "}dummy"), f"{{{ns}}}{tag.split(':')[1]}")
    if attrs:
        for k, v in attrs.items():
            el.set(f"{{{ns}}}{k.split(':')[1]}", str(v))
    return el
