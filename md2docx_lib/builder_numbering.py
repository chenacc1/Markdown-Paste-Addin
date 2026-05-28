"""Auto-numbering tracker for figures, tables, and headings in Word documents."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


"""Auto-numbering tracker for figures, tables, and headings in Word documents."""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class NumberingTracker:
    """Tracks and assigns sequential numbers to figures and tables."""

    def __init__(self):
        self.figure_count = 0
        self.table_count = 0
        self.heading_counts = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}

    def next_figure(self, caption_text: str = "") -> str:
        """Return formatted figure caption, e.g. 'Figure 1: Description'."""
        self.figure_count += 1
        label = f"图 {self.figure_count}"
        if caption_text:
            label = f"图 {self.figure_count}：{caption_text}"
        return label

    def next_table(self, caption_text: str = "") -> str:
        """Return formatted table caption, e.g. 'Table 1: Description'."""
        self.table_count += 1
        label = f"表 {self.table_count}"
        if caption_text:
            label = f"表 {self.table_count}：{caption_text}"
        return label

    def next_heading(self, level: int, text: str) -> str:
        """Track heading and prepend auto-numbering.
        Strips any existing numbering from text to avoid duplication
        (e.g. '2.2 Title' + auto-number '1.2.2' would produce '1.2.2 2.2 Title').
        """
        for l in range(level + 1, 7):
            self.heading_counts[l] = 0
        self.heading_counts[level] += 1

        parts = [str(self.heading_counts[l]) for l in range(1, level + 1)
                 if self.heading_counts.get(l, 0) > 0]

        # Strip existing numbering patterns from the original text
        clean_text = self._strip_existing_numbering(text)

        if parts:
            return ".".join(parts) + " " + clean_text
        return clean_text

    @staticmethod
    def _strip_existing_numbering(text: str) -> str:
        """Remove existing heading numbers like '1.', '1.2', '1.2.3', '1.2.3 ' from text."""
        return re.sub(r"^[\d]+(\.[\d]+)*[\s.、]+", "", text).strip()


def add_figure_caption(doc: Document, tracker: NumberingTracker,
                       caption_text: str = "", alt_text: str = ""):
    """Add a centered figure caption paragraph."""
    caption = tracker.next_figure(caption_text or alt_text)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = "宋体"
    run.bold = False
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    return caption


def add_table_caption(doc: Document, tracker: NumberingTracker,
                      caption_text: str = ""):
    """Add a centered table caption paragraph (placed before the table)."""
    caption = tracker.next_table(caption_text)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    run.font.size = Pt(10)
    run.font.name = "宋体"
    run.bold = True
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return caption
