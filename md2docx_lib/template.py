"""Word template support — load .docx template and append content at marker."""

from pathlib import Path
from docx import Document


def load_template(template_path: str) -> Document:
    """Load a .docx file as template base. Returns a Document object.

    If the file doesn't exist, returns a new empty document.
    """
    if template_path and Path(template_path).exists():
        try:
            return Document(template_path)
        except Exception as e:
            print(f"  [warn] Failed to load template {template_path}: {e}")
    return Document()


def find_marker_paragraph(doc: Document, marker: str = "{{CONTENT}}") -> int | None:
    """Find the paragraph index containing the marker text.

    Returns the paragraph index, or None if not found.
    The content will be inserted at this position, replacing the marker.
    """
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            return i
    return None


def insert_at_marker(doc: Document, marker: str = "{{CONTENT}}"):
    """Find marker paragraph, clear it, and return the paragraph object.

    New content should be added before this paragraph via doc.add_paragraph()
    at the correct position. After insertion, this marker paragraph is deleted.
    """
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text:
            # Clear the marker paragraph text
            para.clear()
            p_element = para._element
            p_element.getparent().remove(p_element)
            return True
    return False
