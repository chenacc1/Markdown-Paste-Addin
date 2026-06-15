"""Word document builder — converts parsed chunks into a .docx file.

Features:
  - Text (headings, lists, blockquotes, code blocks, task lists)
  - Tables with optional auto-captions
  - Mermaid diagrams → embedded PNG with auto-captions
  - Embedded images with auto-captions
  - LaTeX math → OMML native equations
  - Syntax-highlighted code blocks
  - Horizontal rules
  - Progress feedback (tqdm)
  - Error recovery per chunk
"""

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.table import _Cell
from docx.oxml import OxmlElement

from .renderer_mermaid import render_mermaid
from .renderer_code import highlight_code
from .renderer_image import download_image
from .builder_numbering import NumberingTracker
from .builder_toc import add_toc as _add_toc_field
from .inline_processor import process_inline_paragraph, clean_text


# ─── Main entry ─────────────────────────────────────────────────────────────────

class DocumentBuilder:
    """Builds a Word document from parsed chunks."""

    def __init__(self, template_path: str = "", show_progress: bool = True,
                 auto_captions: bool = True, image_max_width: float = 5.5,
                 add_headers: bool = True, add_toc: bool = True,
                 title: str = ""):
        self.template_path = template_path
        self.show_progress = show_progress
        self.auto_captions = auto_captions
        self.image_max_width = image_max_width
        self.add_headers_flag = add_headers
        self.add_toc_flag = add_toc
        self.title = title
        self.numbering = NumberingTracker()
        self.errors: list[str] = []

    def build(self, chunks: list[dict], output_path: str) -> str:
        """Build document from chunks and save. Returns output_path."""
        doc = self._create_document()

        # Auto-detect title from first heading if not provided
        if not self.title:
            for c in chunks:
                if c.get("type") == "heading":
                    self.title = c.get("text", "")
                    break

        # Add TOC at beginning if requested
        if self.add_toc_flag:
            _add_toc_field(doc)

        # Add title if provided
        if self.title:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(self.title)
            run.bold = True
            _set_run_font(run, latin=HEADING_LATIN_FONT, cjk=HEADING_CJK_FONT, size=Pt(22))
            para.paragraph_format.space_after = Pt(18)

        # Progress bar
        iterator = chunks
        if self.show_progress and len(chunks) > 5:
            try:
                from tqdm import tqdm
                iterator = tqdm(chunks, desc="Building Word", unit="chunk")
            except ImportError:
                pass

        for chunk in iterator:
            try:
                self._process_chunk(doc, chunk)
            except Exception as e:
                error_msg = f"[Error processing {chunk.get('type', '?')}]: {e}"
                self.errors.append(error_msg)
                print(f"  [error] {error_msg}")
                # Insert error placeholder
                para = doc.add_paragraph()
                run = para.add_run(f"[转换错误: {e}]")
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        if self.errors:
            print(f"\n  {len(self.errors)} error(s) occurred during conversion.")

        doc.save(output_path)
        return output_path

    def _create_document(self) -> Document:
        from .template import load_template, insert_at_marker
        doc = load_template(self.template_path)

        style = doc.styles["Normal"]
        style.font.name = DEFAULT_LATIN_FONT
        style.font.size = DEFAULT_FONT_SIZE
        style.paragraph_format.line_spacing = 1.5

        # Set CJK font on the Normal style so Chinese text renders in 宋体
        style_rPr = style.element.get_or_add_rPr()
        style_rFonts = style_rPr.find(qn("w:rFonts"))
        if style_rFonts is None:
            style_rFonts = style_rPr.makeelement(qn("w:rFonts"), {})
            style_rPr.insert(0, style_rFonts)
        style_rFonts.set(qn("w:eastAsia"), DEFAULT_CJK_FONT)
        style_rFonts.set(qn("w:ascii"), DEFAULT_LATIN_FONT)
        style_rFonts.set(qn("w:hAnsi"), DEFAULT_LATIN_FONT)

        # Also fix heading styles — built-in defaults use Calibri which has no CJK glyphs
        for level in range(1, 4):
            try:
                h_style = doc.styles[f"Heading {level}"]
                h_style.font.name = HEADING_LATIN_FONT
                h_rPr = h_style.element.get_or_add_rPr()
                h_rFonts = h_rPr.find(qn("w:rFonts"))
                if h_rFonts is None:
                    h_rFonts = h_rPr.makeelement(qn("w:rFonts"), {})
                    h_rPr.insert(0, h_rFonts)
                h_rFonts.set(qn("w:eastAsia"), HEADING_CJK_FONT)
                h_rFonts.set(qn("w:ascii"), HEADING_LATIN_FONT)
                h_rFonts.set(qn("w:hAnsi"), HEADING_LATIN_FONT)
            except KeyError:
                pass  # style not present in template

        return doc

    def _process_chunk(self, doc: Document, chunk: dict):
        t = chunk.get("type", "text")

        if t == "text":
            self._insert_text(doc, chunk)
        elif t == "table":
            self._insert_table(doc, chunk)
        elif t == "mermaid":
            self._insert_mermaid(doc, chunk)
        elif t == "image":
            self._insert_image(doc, chunk)
        elif t == "code":
            self._insert_code(doc, chunk)
        elif t == "math":
            self._insert_math(doc, chunk)
        elif t == "task_list":
            self._insert_task_list(doc, chunk)
        elif t == "blockquote":
            self._insert_blockquote(doc, chunk)
        elif t == "hr":
            self._insert_hr(doc)
        elif t == "heading":
            self._insert_heading(doc, chunk)

    # ── Text ──

    def _insert_text(self, doc: Document, chunk: dict):
        text = chunk.get("text", "")
        for para_text in text.split("\n"):
            para = doc.add_paragraph()

            # Heading pattern: # text
            if para_text.startswith("#"):
                m = __import__('re').match(r"^(#+)", para_text)
                if m:
                    level = min(len(m.group(1)), 9)
                    content = para_text.lstrip("#").strip()
                    numbered = self.numbering.next_heading(level, content)
                    para.style = f"Heading {level}"
                    process_inline_paragraph(para, numbered)
                    continue

            # Unordered list: - / *
            if __import__('re').match(r"^\s*[-*+]\s+", para_text):
                indent = len(para_text) - len(para_text.lstrip())
                level = indent // 2 + 1
                content = __import__('re').sub(r"^\s*[-*+]\s+", "", para_text)
                style_name = "List Bullet" if level == 1 else f"List Bullet {level}"
                try:
                    para.style = style_name
                except KeyError:
                    para.style = "List Bullet"
                process_inline_paragraph(para, content)
                continue

            # Ordered list: 1. / 1) / 1、
            list_match = __import__('re').match(r"^(\s*)(\d+)[.)、]\s*", para_text)
            if list_match:
                indent = len(list_match.group(1))
                level = indent // 2 + 1
                content = __import__('re').sub(r"^\s*\d+[.)、]\s*", "", para_text)
                style_name = "List Number" if level == 1 else f"List Number {level}"
                try:
                    para.style = style_name
                except KeyError:
                    para.style = "List Number"
                process_inline_paragraph(para, content)
                continue

            # Plain text — process inline markdown (**bold**, *italic*, `code`)
            if para_text.strip():
                process_inline_paragraph(para, para_text)
            else:
                para.paragraph_format.space_after = Pt(6)

    # ── Heading ──

    def _insert_heading(self, doc: Document, chunk: dict):
        level = min(chunk.get("level", 1), 9)
        text = chunk.get("text", "")
        numbered = self.numbering.next_heading(level, text)
        para = doc.add_paragraph()
        para.style = f"Heading {level}"
        process_inline_paragraph(para, numbered)

    # ── Table ──

    def _insert_table(self, doc: Document, chunk: dict):
        headers = chunk["headers"]
        alignments = chunk.get("alignments", ["left"] * len(headers))
        rows = chunk["rows"]
        total_rows = len(rows) + 1
        total_cols = len(headers)

        # Auto-caption
        if self.auto_captions and chunk.get("caption"):
            from .builder_numbering import add_table_caption
            add_table_caption(doc, self.numbering, chunk.get("caption", ""))

        table = doc.add_table(rows=total_rows, cols=total_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for col_idx, header_text in enumerate(headers):
            cell = table.cell(0, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = _wd_align(alignments[col_idx] if col_idx < len(alignments) else "left")
            run = p.add_run(header_text)
            run.bold = True
            _set_run_font(run, latin=DEFAULT_LATIN_FONT, cjk=DEFAULT_CJK_FONT, size=Pt(10))
            _set_cell_shading(cell, "D9E2F3")

        # Data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < total_cols:
                    cell = table.cell(row_idx + 1, col_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = _wd_align(alignments[col_idx] if col_idx < len(alignments) else "left")
                    run = p.add_run(cell_text)
                    _set_run_font(run, latin=DEFAULT_LATIN_FONT, cjk=DEFAULT_CJK_FONT, size=Pt(10))

        doc.add_paragraph()

    # ── Mermaid ──

    def _insert_mermaid(self, doc: Document, chunk: dict):
        png_bytes = render_mermaid(chunk["code"])
        if png_bytes:
            self._embed_image_bytes(doc, png_bytes)
            # Auto-caption
            if self.auto_captions:
                from .builder_numbering import add_figure_caption
                add_figure_caption(doc, self.numbering,
                                   chunk.get("caption", ""))
        else:
            para = doc.add_paragraph()
            run = para.add_run("[流程图无法渲染]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.add_paragraph()

    # ── Image ──

    def _insert_image(self, doc: Document, chunk: dict):
        src = chunk.get("src", "")
        alt = chunk.get("alt", "")
        print(f"  Downloading image: {src[:80]}...")
        img_bytes = download_image(src)
        if img_bytes:
            self._embed_image_bytes(doc, img_bytes)
            if self.auto_captions:
                from .builder_numbering import add_figure_caption
                add_figure_caption(doc, self.numbering, alt, alt)
            print("  OK")
        else:
            para = doc.add_paragraph()
            run = para.add_run(f"[Image: {alt or src}]")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            print("  FAILED")
        doc.add_paragraph()

    # ── Code block with syntax highlighting ──

    def _insert_code(self, doc: Document, chunk: dict):
        code = chunk.get("code", "")
        language = chunk.get("language", "")

        # Syntax-highlighted lines
        highlighted = highlight_code(code, language)

        if highlighted:
            for line_tokens in highlighted:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0

                # Light gray background via shading
                _set_paragraph_shading(para, "F5F5F5")

                for token in line_tokens:
                    run = para.add_run(token["text"])
                    _set_run_font(run, latin=CODE_FONT, cjk=CODE_FONT, size=CODE_FONT_SIZE)
                    style = token.get("style", {})
                    if "color" in style:
                        run.font.color.rgb = style["color"]
                    if style.get("bold"):
                        run.bold = True
                    if style.get("italic"):
                        run.italic = True

        # Add spacing after code block
        doc.add_paragraph()

    # ── Math ──

    def _insert_math(self, doc: Document, chunk: dict):
        latex = chunk.get("latex", "")
        display = chunk.get("display", False)

        from .parser_math import latex_to_omml, _omml_wrap

        try:
            inner_omml = latex_to_omml(latex, display)
            wrapped_xml = _omml_wrap(inner_omml, display)
            if display:
                # Display math: centered, standalone paragraph
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _insert_omml_element(para, wrapped_xml)
            else:
                # Inline math: append to the last paragraph if it exists
                # and is not empty/structured, otherwise create new paragraph
                if doc.paragraphs:
                    last_para = doc.paragraphs[-1]
                    last_text = last_para.text.strip()
                    # Only append to existing paragraph if it has plain text
                    # and doesn't already contain math elements
                    has_math = (
                        last_para._element.findall(qn("m:oMath")) or
                        last_para._element.findall(qn("m:oMathPara")))
                    if last_text and not has_math:
                        _insert_omml_element(last_para, wrapped_xml)
                    else:
                        para = doc.add_paragraph()
                        _insert_omml_element(para, wrapped_xml)
                else:
                    para = doc.add_paragraph()
                    _insert_omml_element(para, wrapped_xml)
        except Exception as e:
            para = doc.add_paragraph()
            run = para.add_run(f"${latex}$")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Task list ──

    def _insert_task_list(self, doc: Document, chunk: dict):
        items = chunk.get("items", [])
        for item in items:
            para = doc.add_paragraph()
            checkbox = "☒" if item.get("checked") else "☐"
            run = para.add_run(f"{checkbox}  {item.get('text', '')}")
            _set_run_font(run)

    # ── Blockquote ──

    def _insert_blockquote(self, doc: Document, chunk: dict):
        text = chunk.get("text", "")
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.5)
        para.paragraph_format.right_indent = Cm(0.5)
        _set_paragraph_shading(para, "F0F0F0")
        # Process inline markdown first, then style all runs
        process_inline_paragraph(para, text)
        for run in para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_run_font(run, size=Pt(10.5))

    # ── Horizontal rule ──

    def _insert_hr(self, doc: Document):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        # Insert a bottom border on the paragraph to simulate a horizontal rule
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Image embedding ──

    def _embed_image_bytes(self, doc: Document, img_bytes: bytes,
                           max_width: float = None):
        if max_width is None:
            max_width = self.image_max_width
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            f.write(img_bytes)
            tmp_path = f.name

        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp_path, width=Inches(max_width))
        finally:
            os.unlink(tmp_path)


# ─── Convenience function ───────────────────────────────────────────────────────

def build_docx(chunks: list[dict], output_path: str, **kwargs) -> str:
    """Build Word document from chunks. One-liner convenience."""
    builder = DocumentBuilder(**kwargs)
    return builder.build(chunks, output_path)


# ─── Font helper — ensures both Latin and CJK fonts are set on every run ────────

DEFAULT_LATIN_FONT = "Times New Roman"
DEFAULT_CJK_FONT = "宋体"
DEFAULT_FONT_SIZE = Pt(12)
HEADING_CJK_FONT = "黑体"
HEADING_LATIN_FONT = "Arial"
CODE_FONT = "Consolas"
CODE_FONT_SIZE = Pt(9)


def _set_run_font(run, latin: str = None, cjk: str = None, size=None):
    """Set both Latin (.font.name) and CJK (rFonts.eastAsia) fonts on a run.

    Calling run.font.name alone only sets the Latin font. Chinese glyphs need
    the East-Asian attribute on the rFonts element.
    """
    if latin is None:
        latin = DEFAULT_LATIN_FONT
    if cjk is None:
        cjk = DEFAULT_CJK_FONT
    if size is None:
        size = DEFAULT_FONT_SIZE

    run.font.name = latin
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def _set_style_default_font(doc: Document, latin: str = None, cjk: str = None, size=None):
    """Set both Latin and CJK default fonts on the Normal style."""
    if latin is None:
        latin = DEFAULT_LATIN_FONT
    if cjk is None:
        cjk = DEFAULT_CJK_FONT
    if size is None:
        size = DEFAULT_FONT_SIZE

    style = doc.styles["Normal"]
    style.font.name = latin
    style.font.size = size

    # Set East-Asian font on the style's rPr
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


# ─── Helpers ────────────────────────────────────────────────────────────────────

def _wd_align(align: str):
    if align == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _set_cell_shading(cell: _Cell, color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    tcPr.append(shading)


def _set_paragraph_shading(para, color: str):
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    pPr.append(shading)


def _insert_omml_element(para, omml_xml: str):
    """Insert OMML equation XML into a paragraph.

    The omml_xml should be a complete, well-formed OMML element
    (e.g. <m:oMathPara> or <m:oMath>) with namespace declarations.

    Both display math (<m:oMathPara>) and inline math (<m:oMath>)
    are appended directly as children of <w:p>. This is the correct
    OMML structure — math elements are siblings of <w:r> within <w:p>.
    """
    from lxml import etree as _etree
    try:
        omml_el = _etree.fromstring(omml_xml.encode("utf-8"))
        # Both <m:oMathPara> and <m:oMath> are direct children of <w:p>
        para._element.append(omml_el)
    except Exception as e:
        # Fallback: insert as plain text
        run = para.add_run("[公式]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
