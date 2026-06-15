"""Word document format unifier — auto-detect heading/caption/body types
and apply consistent formatting.

Enhanced from format_docx.py with:
  - Multi-language detection (Chinese + English)
  - Exported reusable functions (detect_paragraph_type, apply_format, format_document)
  - Per-section styling for headers/footers
  - Better heading numbering detection
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# ═══════════════════════════════════════════════════════════════════════════════
# Format presets — customizable
# ═══════════════════════════════════════════════════════════════════════════════

FORMAT = {
    "一级标题": {
        "font_name": "黑体",
        "font_size": Pt(16),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(18),
        "space_after": Pt(12),
        "line_spacing": 1.5,
        "outline_level": 1,
    },
    "二级标题": {
        "font_name": "黑体",
        "font_size": Pt(14),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(14),
        "space_after": Pt(8),
        "line_spacing": 1.5,
        "outline_level": 2,
    },
    "三级标题": {
        "font_name": "黑体",
        "font_size": Pt(12),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(10),
        "space_after": Pt(6),
        "line_spacing": 1.5,
        "outline_level": 3,
    },
    "图题": {
        "font_name": "宋体",
        "font_size": Pt(10),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(6),
        "space_after": Pt(12),
        "line_spacing": 1.5,
    },
    "表题": {
        "font_name": "宋体",
        "font_size": Pt(10),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(12),
        "space_after": Pt(6),
        "line_spacing": 1.5,
    },
    "正文": {
        "font_name": "宋体",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "line_spacing": 1.5,
        "first_line_indent": Cm(0.74),
    },
}

_EN_FORMAT = {
    "一级标题": {
        "font_name": "Arial",
        "font_size": Pt(16),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(18),
        "space_after": Pt(12),
        "line_spacing": 1.15,
        "outline_level": 1,
    },
    "二级标题": {
        "font_name": "Arial",
        "font_size": Pt(14),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(14),
        "space_after": Pt(8),
        "line_spacing": 1.15,
        "outline_level": 2,
    },
    "三级标题": {
        "font_name": "Arial",
        "font_size": Pt(12),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(10),
        "space_after": Pt(6),
        "line_spacing": 1.15,
        "outline_level": 3,
    },
    "图题": {
        "font_name": "Times New Roman",
        "font_size": Pt(10),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(6),
        "space_after": Pt(12),
        "line_spacing": 1.15,
    },
    "表题": {
        "font_name": "Times New Roman",
        "font_size": Pt(10),
        "bold": True,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(12),
        "space_after": Pt(6),
        "line_spacing": 1.15,
    },
    "正文": {
        "font_name": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "color": RGBColor(0, 0, 0),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(0),
        "space_after": Pt(0),
        "line_spacing": 1.15,
        "first_line_indent": Cm(1.27),
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# Paragraph type detection
# ═══════════════════════════════════════════════════════════════════════════════

def detect_paragraph_type(paragraph, lang: str = "cn") -> str | None:
    """Detect the semantic type of a paragraph.

    Args:
        paragraph: python-docx Paragraph object
        lang: 'cn' for Chinese patterns, 'en' for English
    """
    text = paragraph.text.strip()
    if not text:
        return "空行"

    # Already has a heading style
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "heading 1" in style_name or "heading1" in style_name:
        return "一级标题"
    if "heading 2" in style_name or "heading2" in style_name:
        return "二级标题"
    if "heading 3" in style_name or "heading3" in style_name:
        return "三级标题"

    if lang == "cn":
        return _detect_cn(text, paragraph)
    else:
        return _detect_en(text, paragraph)


def _detect_cn(text: str, paragraph) -> str:
    if re.match(r"^(图|Fig\.?|Figure)\s*\d+", text, re.IGNORECASE):
        return "图题"
    if re.match(r"^(表|Table)\s*\d+", text, re.IGNORECASE):
        return "表题"
    if re.match(r"^(第[一二三四五六七八九十\d]+章|第[一二三四五六七八九十\d]+节|[一二三四五六七八九十]、)", text):
        return "一级标题"
    if re.match(r"^[\(\（][一二三四五六七八九十\d]+[\)\）]", text):
        return "二级标题"
    if re.match(r"^\d+[\.\、]\s*\S", text) and len(text) < 60:
        return "三级标题"

    # Bold short text → heading
    if len(text) < 40 and paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.bold and first_run.font.size:
            if first_run.font.size >= Pt(16):
                return "一级标题"
            elif first_run.font.size >= Pt(14):
                return "二级标题"

    return "正文"


def _detect_en(text: str, paragraph) -> str:
    if re.match(r"^(Figure|Fig\.?)\s*\d+", text, re.IGNORECASE):
        return "图题"
    if re.match(r"^Table\s*\d+", text, re.IGNORECASE):
        return "表题"
    if re.match(r"^(Chapter|Section)\s+\d+", text, re.IGNORECASE):
        return "一级标题"
    if re.match(r"^\d+\.\d+\s", text) and len(text) < 80:
        return "二级标题"
    if re.match(r"^\d+\.\s", text) and len(text) < 60:
        return "三级标题"

    if len(text) < 40 and paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.bold and first_run.font.size:
            if first_run.font.size >= Pt(16):
                return "一级标题"
            elif first_run.font.size >= Pt(14):
                return "二级标题"

    return "正文"


# ═══════════════════════════════════════════════════════════════════════════════
# Format application
# ═══════════════════════════════════════════════════════════════════════════════

def apply_format(paragraph, style_name: str, format_dict: dict = None):
    """Apply predefined formatting to a paragraph."""
    fmt = format_dict or FORMAT.get(style_name)
    if not fmt:
        return

    pf = paragraph.paragraph_format

    if "space_before" in fmt:
        pf.space_before = fmt["space_before"]
    if "space_after" in fmt:
        pf.space_after = fmt["space_after"]
    if "line_spacing" in fmt:
        pf.line_spacing = fmt["line_spacing"]
    if "alignment" in fmt:
        pf.alignment = fmt["alignment"]
    if "first_line_indent" in fmt:
        pf.first_line_indent = fmt["first_line_indent"]

    # Outline level (for navigation pane/TOC)
    if "outline_level" in fmt:
        pPr = paragraph._element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:outlineLvl")):
            pPr.remove(existing)
        outline_lvl = pPr.makeelement(qn("w:outlineLvl"), {})
        outline_lvl.set(qn("w:val"), str(fmt["outline_level"]))
        pPr.append(outline_lvl)

    # Apply font to all runs
    for run in paragraph.runs:
        _apply_run_font(run, fmt)

    # If no runs but has text, add a run
    if not paragraph.runs and paragraph.text.strip():
        run = paragraph.add_run(paragraph.text)
        _apply_run_font(run, fmt)


def _apply_run_font(run, fmt: dict):
    if "font_name" in fmt:
        run.font.name = fmt["font_name"]
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), fmt["font_name"])
    if "font_size" in fmt:
        run.font.size = fmt["font_size"]
    if "bold" in fmt:
        run.font.bold = fmt["bold"]
    if "color" in fmt:
        run.font.color.rgb = fmt["color"]


def set_default_font(doc: Document, font_name: str = "宋体", font_size: Pt = Pt(12)):
    """Set both Latin and CJK default fonts on the Normal style."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = font_size
    style.paragraph_format.line_spacing = 1.5

    # Set East-Asian font on the style element so Chinese text renders correctly
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════════════════════════════════════════
# Main formatting function
# ═══════════════════════════════════════════════════════════════════════════════

def format_document(doc: Document, lang: str = "cn", format_dict: dict = None) -> dict:
    """Format all paragraphs in a document. Returns statistics."""
    fmt = format_dict or FORMAT
    stats = {}

    set_default_font(doc)

    for paragraph in doc.paragraphs:
        ptype = detect_paragraph_type(paragraph, lang)
        if ptype == "空行":
            stats["空行"] = stats.get("空行", 0) + 1
            continue
        apply_format(paragraph, ptype, fmt)
        stats[ptype] = stats.get(ptype, 0) + 1

    # Format paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_format(paragraph, "正文", fmt)

    # Format headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            if para.text.strip():
                apply_format(para, "正文", fmt)

    return stats
